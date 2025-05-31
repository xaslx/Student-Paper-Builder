import os
from docx import Document
from docx.shared import Pt, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate, InlineImage
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
from PIL import Image, UnidentifiedImageError


class DocxCreator:

    def __init__(self, template_path: str):
        self.template_path = template_path

    def _process_image(self, doc_tpl: DocxTemplate, image_path: str) -> dict:

        if not os.path.exists(image_path):
            return None

        try:
            with Image.open(image_path) as img:

                if img.format not in ('PNG', 'JPEG', 'WEBP', 'JPG'):
                    return None


                if img.format == 'WEBP':
                    try:
                        png_path = os.path.splitext(image_path)[0] + '.png'
                        img.save(png_path, 'PNG')
                        image_path = png_path
                    except Exception:
                        return None


                max_width_mm = 167
                max_height_mm = 100
                

                max_width_pt = max_width_mm * 2.83465
                max_height_pt = max_height_mm * 2.83465


                width, height = img.size
                

                width_ratio = max_width_pt / width
                height_ratio = max_height_pt / height
                

                ratio = min(width_ratio, height_ratio)
                

                new_width = int(width * ratio)
                new_height = int(height * ratio)
                

                new_width_mm = new_width / 2.83465
                new_height_mm = new_height / 2.83465

                return {
                    'path': image_path,
                    'image': InlineImage(
                        doc_tpl, 
                        image_path, 
                        width=Mm(new_width_mm), 
                        height=Mm(new_height_mm))
                }

        except UnidentifiedImageError:
            return None
        except Exception:
            return None


    def _create_element(self, name: str) -> OxmlElement:
        return OxmlElement(name)

    def _create_attribute(self, element: OxmlElement, name: str, value: str) -> None:
        element.set(qn(name), value)

    def _add_page_number(self, run) -> None:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        fldChar1 = self._create_element('w:fldChar')
        self._create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self._create_element('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = self._create_element('w:fldChar')
        self._create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _add_footer(self, doc: Document, city: str, year: int) -> None:
        section = doc.sections[0]
        section.footer_distance = Pt(10)


        first_page_footer = section.first_page_footer
        paragraph = first_page_footer.paragraphs[0] if first_page_footer.paragraphs else first_page_footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        

        run = paragraph.add_run(f'{city}, {year}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.superscript = False
        

        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(0)

        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        self._add_page_number(run)

    def fill_template_with_breaks(self, context: dict, output_path: str) -> None:
        doc_tpl = DocxTemplate(self.template_path)


        appendices = context.get('appendices', [])

        valid_appendices = []
        
        for appendix in appendices:
            if not appendix.get('path'):
                continue
                
            processed = self._process_image(doc_tpl, appendix['path'])
            if processed:
                valid_appendices.append({**appendix, **processed})

        context['appendices'] = valid_appendices


        doc_tpl.render(context=context)


        temp_docx_path = f'src/presentation/static/docx/{output_path}_temp.docx'
        doc_tpl.save(temp_docx_path)


        final_doc = Document(temp_docx_path)
        final_doc.sections[0].different_first_page_header_footer = True

        title_page = context.get('title_page', {})
        city = title_page.get('city', 'Город')
        year = title_page.get('year', 2025)
        self._add_footer(final_doc, city, year)


        output_docx_path = f'src/presentation/static/docx/{output_path}.docx'
        final_doc.save(output_docx_path)
        try:
            os.remove(temp_docx_path)
        except OSError as e:
            print(f'Ошибка при удалении временного файла: {e}')

        try:
            self._convert_to_pdf(
                docx_path=output_docx_path,
                pdf_path=f'src/presentation/static/pdf/{output_path}.pdf',
            )
        except subprocess.CalledProcessError as e:
            print(f'Ошибка конвертации: {e}')
            raise subprocess.CalledProcessError()

    def _convert_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', '/'.join(pdf_path.split('/')[:-1]),
            docx_path
        ], check=True)